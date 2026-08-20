from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header / Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run("MAX BALLARD")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run("Kansas City, MO / Open to Remote | ballardcmax@gmail.com | github.com/Champlin92")
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph(
        "IT Manager specializing in Infrastructure Operations, Enterprise Batch/EFT, and Proactive Observability. "
        "Proven track record of scaling high-availability support for 3,600+ North American retail stores, "
        "protecting $420k/hr in revenue, and driving 70% faster development through strategic AI enablement."
    )

    doc.add_heading('QUANTIFIABLE IMPACT & METRICS', level=1)
    metrics = [
        ("Revenue Protected / Hr: ", "$420k - Safeguarded via custom Python observability pipelines preventing Sev-1 outages across 3,600+ stores."),
        ("Labor Hours Saved Annually: ", "1,300+ - Reclaimed by bypassing L1/L2 and routing 30+ daily POS errors directly to L3 via ServiceNow."),
        ("Dev Acceleration: ", "70% - Boosted script development speed by integrating Claude AI securely into engineering workflows."),
        ("Team Certification: ", "100% - Led by example, achieving 100% Python certification rates across the automation engineering team.")
    ]
    for bold_text, normal_text in metrics:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)

    # Role 1
    p_role1 = doc.add_paragraph()
    p_role1.add_run("IT Manager, InfraOps & Automation").bold = True
    p_role1.add_run("\nSally Beauty Holdings | Denton, TX / Remote | March 2024 – Present").italic = True
    
    bullets1 = [
        ("Enterprise Scale: ", "Direct management of Automation, App Mgmt, and Batch operations for 3,600+ retail stores. Manage critical vendor relations (Fortra, Vinzant, PagerDuty)."),
        ("AI Enablement & Strategy: ", "Accelerated script development by 70% and streamlined reviews by integrating Claude AI. Directed enterprise architecture strategy, steering stakeholders toward low-cost deterministic automation over high-cost LLM deployments."),
        ("Proactive Observability ($420k/hr impact): ", "Architected a custom Python/Batch observability pipeline that auto-detects backend WebLogic failures, instantly routes alerts to L3 engineering, and autonomously suppresses downstream POS errors, completely preempting Sev-1 call waves."),
        ("Call Volume Optimization: ", "Built a self-healing pipeline integrating directly with ServiceNow to bypass L1/L2 triage. Automatically resolved/routed 30+ daily POS errors, saving over 1,300 labor hours annually and mitigating up to 375 hours of acute downtime during severity events."),
        ("Hands-On Leadership: ", "Prevented a company-wide onboarding/offboarding outage by diagnosing a breaking production script within a 60-minute SLA, collaboratively walking the engineering team through root-cause analysis and remediation. Led by example, achieving 100% team Python certification rates by taking the courses alongside them.")
    ]
    for bold_text, normal_text in bullets1:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    # Role 2
    p_role2 = doc.add_paragraph()
    p_role2.add_run("Supervisor, Middleware Team").bold = True
    p_role2.add_run("\nSally Beauty Holdings | Denton, TX | March 2023 – March 2024").italic = True
    
    bullets2 = [
        "Managed daily operations, workload prioritization, and project delegation for the Middleware engineering team within Infrastructure Operations.",
        "Acted as primary escalation lead for complex engineering-level incidents across enterprise backend systems, integration tiers, and retail POS environments.",
        "Partnered with application owners and platform leads to uphold system reliability, SLA adherence, and change governance."
    ]
    for bullet in bullets2:
        doc.add_paragraph(bullet, style='List Bullet')

    # Role 3
    p_role3 = doc.add_paragraph()
    p_role3.add_run("Middleware Systems Engineer").bold = True
    p_role3.add_run("\nSally Beauty Holdings | Denton, TX | February 2021 – March 2023").italic = True
    
    bullets3 = [
        ("Critical Enterprise Deployment: ", "Stepped into a leadership vacuum following the sudden departure of the team lead to engineer the deployment of a 6-node Oracle WebLogic cluster serving as the backbone for all POS systems. Resolved severe clustering failures during a 24-hour continuous window, preventing the delay of a massive enterprise POS upgrade."),
        ("De-Siloed Critical EFT Operations: ", "Transitioned critical Globalscape EFT workload from a single point of failure into a resilient team discipline, writing comprehensive documentation and leading cross-training."),
        ("Knowledgebase Architecture: ", "Authored 55% of all technical knowledgebase articles and runbooks on a team of six engineers."),
        ("Service Desk Tooling: ", "Developed, maintained, and upgraded specialized troubleshooting toolsets empowering Service Desk Leads to resolve POS issues independently.")
    ]
    for bold_text, normal_text in bullets3:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    # Role 4
    p_role4 = doc.add_paragraph()
    p_role4.add_run("Service Desk Lead & Analyst").bold = True
    p_role4.add_run("\nSally Beauty Holdings | Denton, TX | February 2019 – February 2021").italic = True
    
    bullets4 = [
        ("Lead (2020–2021): ", "Led a multidisciplinary technical crew through COVID-era operations, serving as the high-severity incident commander. Authored foundational helpdesk training guides and troubleshooting standards."),
        ("Analyst (2019–2020): ", "Provided high-velocity rotational on-call and after-hours support, diagnosing and resolving complex hardware, software, and network connectivity incidents.")
    ]
    for bold_text, normal_text in bullets4:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    doc.add_heading('TECHNICAL SKILLS & COMPETENCIES', level=1)
    skills = [
        ("Scripting & Automation: ", "Python, Bash / POSIX Shell, PowerShell, REST API Integrations, Oracle SQL"),
        ("Infrastructure & Middleware: ", "Globalscape EFT & Batch, Oracle WebLogic, Red Hat Linux (RHEL), Apache HTTP, Windows Server, Azure"),
        ("Observability & Operations: ", "Dynatrace, Rundeck Orchestration, Oracle Xstore POS Diagnostics, Batch Job Scheduling")
    ]
    for bold_text, normal_text in skills:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    doc.add_heading('EDUCATION & CERTIFICATIONS', level=1)
    edu = [
        "Google AI Professional Certification",
        "Google IT Support Professional Certification",
        "ITIL 4 Foundation (In Progress)",
        "Undergraduate Studies in Anthropology – University of Wyoming",
        "Undergraduate Studies in Biology – University of Missouri–Kansas City"
    ]
    for e in edu:
        doc.add_paragraph(e, style='List Bullet')

    doc.save("Max_Ballard_Resume_Updated.docx")

if __name__ == "__main__":
    main()
